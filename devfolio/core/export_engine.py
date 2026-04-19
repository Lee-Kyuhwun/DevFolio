"""Markdown / PDF / DOCX / HTML / CSV 내보내기 엔진."""

from __future__ import annotations

import csv
import html as html_mod
import io
import re
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from devfolio.models.project import Project

from devfolio.core.storage import EXPORTS_DIR
from devfolio.exceptions import DevfolioExportError
from devfolio.log import get_logger

logger = get_logger(__name__)

_PDF_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700&display=swap');
body {
    font-family: 'Noto Sans KR', 'Malgun Gothic', 'Apple SD Gothic Neo', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.7;
    color: #2c2c2c;
    margin: 0;
    padding: 0;
}
.container { max-width: 820px; margin: 0 auto; padding: 32px; }
h1 { font-size: 22pt; color: #1a1a2e; border-bottom: 3px solid #1a1a2e; padding-bottom: 8px; }
h2 { font-size: 16pt; color: #16213e; border-bottom: 1px solid #ccc; padding-bottom: 4px; margin-top: 28px; }
h3 { font-size: 13pt; color: #0f3460; margin-top: 20px; }
h4 { font-size: 11pt; color: #533483; margin-top: 14px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; }
th, td { border: 1px solid #ddd; padding: 7px 12px; text-align: left; font-size: 10pt; }
th { background-color: #f0f0f0; font-weight: bold; }
ul { padding-left: 22px; }
li { margin: 3px 0; }
hr { border: none; border-top: 1px solid #ddd; margin: 24px 0; }
code { background: #f5f5f5; padding: 2px 5px; border-radius: 3px; font-family: monospace; font-size: 10pt; }
em { color: #555; }
"""

_HTML_EXTRA_CSS = """
body { background-color: #f0f0f0; }
.container {
    background-color: #fff;
    border-radius: 8px;
    box-shadow: 0 4px 16px rgba(0,0,0,0.08);
    padding: 48px;
    margin: 24px auto;
}
a { color: #0f3460; text-decoration: none; }
a:hover { text-decoration: underline; }
.mermaid {
    margin: 18px 0;
    padding: 12px;
    border-radius: 14px;
    background: #faf7f2;
    overflow-x: auto;
}
"""

_MERMAID_BOOTSTRAP = """
<script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.min.js"></script>
<script>
  document.addEventListener("DOMContentLoaded", async () => {
    if (typeof mermaid === "undefined") return;
    mermaid.initialize({ startOnLoad: false, securityLevel: "loose" });
    const blocks = document.querySelectorAll("pre code.language-mermaid");
    let index = 0;
    for (const code of blocks) {
      const pre = code.closest("pre");
      if (!pre) continue;
      const container = document.createElement("div");
      container.className = "mermaid";
      container.id = `devfolio-mermaid-${index++}`;
      container.textContent = code.textContent;
      pre.replaceWith(container);
    }
    try {
      await mermaid.run({ querySelector: ".mermaid" });
    } catch (error) {
      console.warn("Mermaid render failed", error);
    }
  });
</script>
"""


def _sanitize_filename(filename: str) -> str:
    """파일명에서 디렉터리 구분자와 위험 문자를 제거."""
    filename = filename.replace("/", "_").replace("\\", "_")
    filename = filename.replace("..", "_")
    return filename


def _validate_output_path(path: Path) -> Path:
    """출력 경로가 안전한지 검증 (path traversal 방지)."""
    resolved = path.resolve()
    allowed_roots = [
        Path.home().resolve(),
        Path.cwd().resolve(),
        Path(tempfile.gettempdir()).resolve(),
    ]
    if not any(str(resolved).startswith(str(root)) for root in allowed_roots):
        raise DevfolioExportError(
            f"출력 경로가 허용 범위를 벗어났습니다: {path}",
            hint="홈 디렉터리, 현재 작업 디렉터리, 또는 시스템 임시 디렉터리 하위 경로를 지정하세요.",
        )
    return resolved


class ExportEngine:
    def __init__(self):
        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def build_html_document(html_body: str, title: str = "DevFolio") -> str:
        return f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html_mod.escape(title)}</title>
<style>
{_PDF_CSS}
{_HTML_EXTRA_CSS}
</style>
</head>
<body>
<div class="container">
{html_body}
</div>
{_MERMAID_BOOTSTRAP}
</body>
</html>"""

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def export_markdown(self, content: str, filename: str, output_dir: Path | None = None) -> Path:
        filename = _sanitize_filename(filename)
        if not filename.endswith(".md"):
            filename = filename + ".md"
        base = output_dir or EXPORTS_DIR
        base.mkdir(parents=True, exist_ok=True)
        output_path = _validate_output_path(base / filename)
        output_path.write_text(content, encoding="utf-8")
        return output_path

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def export_pdf(self, content: str, filename: str, output_dir: Path | None = None) -> Path:
        filename = _sanitize_filename(filename)
        if not filename.endswith(".pdf"):
            filename = filename + ".pdf"
        base = output_dir or EXPORTS_DIR
        base.mkdir(parents=True, exist_ok=True)
        output_path = _validate_output_path(base / filename)

        try:
            from weasyprint import CSS, HTML
        except ImportError:
            raise RuntimeError(
                "WeasyPrint가 설치되지 않았습니다.\n"
                "`pip install weasyprint`으로 설치하세요.\n"
                "시스템 패키지(GTK, Pango 등)가 추가로 필요할 수 있습니다."
            )

        html_body = self._md_to_html_body(content)
        full_html = f"""<!DOCTYPE html>
<html lang="ko"><head>
<meta charset="UTF-8">
<style>{_PDF_CSS}</style>
</head><body>
<div class="container">{html_body}</div>
</body></html>"""

        HTML(string=full_html).write_pdf(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def export_docx(self, content: str, filename: str, output_dir: Path | None = None) -> Path:
        filename = _sanitize_filename(filename)
        if not filename.endswith(".docx"):
            filename = filename + ".docx"
        base = output_dir or EXPORTS_DIR
        base.mkdir(parents=True, exist_ok=True)
        output_path = _validate_output_path(base / filename)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise RuntimeError(
                "python-docx가 설치되지 않았습니다.\n"
                "`pip install python-docx`으로 설치하세요."
            )

        doc = Document()
        self._md_to_docx(content, doc)
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # HTML
    # ------------------------------------------------------------------

    def export_html(self, content: str, filename: str, output_dir: Path | None = None) -> Path:
        filename = _sanitize_filename(filename)
        if not filename.endswith(".html"):
            filename = filename + ".html"
        base = output_dir or EXPORTS_DIR
        base.mkdir(parents=True, exist_ok=True)
        output_path = _validate_output_path(base / filename)

        html_body = self._md_to_html_body(content)
        full_html = self.build_html_document(html_body)
        output_path.write_text(full_html, encoding="utf-8")
        return output_path

    # ------------------------------------------------------------------
    # CSV
    # ------------------------------------------------------------------

    def export_csv(self, projects: "list[Project]", filename: str, output_dir: Path | None = None) -> Path:
        """프로젝트 목록을 CSV로 내보낸다.

        각 행은 프로젝트 하나를 나타내며, 태스크 목록은 세미콜론으로 구분한다.

        Args:
            projects: 내보낼 프로젝트 목록
            filename: 출력 파일명 (확장자 없어도 자동 추가)

        Returns:
            생성된 CSV 파일의 경로
        """
        filename = _sanitize_filename(filename)
        if not filename.endswith(".csv"):
            filename = filename + ".csv"
        base = output_dir or EXPORTS_DIR
        base.mkdir(parents=True, exist_ok=True)
        output_path = _validate_output_path(base / filename)

        fieldnames = [
            "id", "name", "type", "status", "organization",
            "period", "role", "team_size", "tech_stack",
            "summary", "tags", "task_count", "tasks",
        ]

        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()

        for p in projects:
            task_names = "; ".join(t.name for t in p.tasks)
            writer.writerow({
                "id": p.id,
                "name": p.name,
                "type": p.type,
                "status": p.status,
                "organization": p.organization,
                "period": p.period.display(),
                "role": p.role,
                "team_size": p.team_size,
                "tech_stack": "; ".join(p.tech_stack),
                "summary": p.summary,
                "tags": "; ".join(p.tags),
                "task_count": len(p.tasks),
                "tasks": task_names,
            })

        output_path.write_text(buf.getvalue(), encoding="utf-8-sig")  # utf-8-sig: Excel 호환
        logger.debug("CSV export: %s (%d rows)", output_path, len(projects))
        return output_path

    # ------------------------------------------------------------------
    # 파일을 지정 경로로 복사
    # ------------------------------------------------------------------

    def copy_to(self, source: Path, destination: Path) -> Path:
        import shutil
        _validate_output_path(destination)
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
        return destination

    # ------------------------------------------------------------------
    # 내부: Markdown → HTML body
    # ------------------------------------------------------------------

    def _md_to_html_body(self, content: str) -> str:
        try:
            import markdown as md_lib
            return md_lib.markdown(
                content,
                extensions=["tables", "fenced_code", "nl2br"],
            )
        except ImportError:
            return self._simple_md_to_html(content)

    def _simple_md_to_html(self, content: str) -> str:
        """최소 Markdown → HTML 변환 (markdown 패키지 없을 때 폴백)."""
        lines = content.split("\n")
        out: list[str] = []
        in_ul = False

        def close_list():
            nonlocal in_ul
            if in_ul:
                out.append("</ul>")
                in_ul = False

        def inline(text: str) -> str:
            text = html_mod.escape(text)
            text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
            text = re.sub(r"\*(.*?)\*", r"<em>\1</em>", text)
            text = re.sub(r"`(.*?)`", r"<code>\1</code>", text)
            return text

        for line in lines:
            if line.startswith("#### "):
                close_list(); out.append(f"<h4>{inline(line[5:])}</h4>")
            elif line.startswith("### "):
                close_list(); out.append(f"<h3>{inline(line[4:])}</h3>")
            elif line.startswith("## "):
                close_list(); out.append(f"<h2>{inline(line[3:])}</h2>")
            elif line.startswith("# "):
                close_list(); out.append(f"<h1>{inline(line[2:])}</h1>")
            elif line.startswith("- ") or line.startswith("* "):
                if not in_ul:
                    out.append("<ul>"); in_ul = True
                out.append(f"<li>{inline(line[2:])}</li>")
            elif line.strip() == "---":
                close_list(); out.append("<hr>")
            elif not line.strip():
                close_list()
            else:
                close_list(); out.append(f"<p>{inline(line)}</p>")

        close_list()
        return "\n".join(out)

    # ------------------------------------------------------------------
    # 내부: Markdown → python-docx
    # ------------------------------------------------------------------

    def _md_to_docx(self, content: str, doc):
        for line in content.split("\n"):
            if line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                self._add_inline(p, line[2:])
            elif line.strip() == "---":
                doc.add_paragraph("─" * 50)
            elif line.strip():
                p = doc.add_paragraph()
                self._add_inline(p, line)

    def _add_inline(self, paragraph, text: str):
        """굵기/기울임 인라인 서식 적용."""
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)
